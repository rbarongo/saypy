from docx import Document
from docx.shared import Pt

DOC_PATH = "Documentation.docx"

doc = Document(DOC_PATH)

doc.add_paragraph("")
heading = doc.add_paragraph("LOCAL COLLECTION CODES - LOCAL ADMIN VIEW")
if heading.runs:
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(12)

lines = [
    "In Local Collection Codes:",
    "These codes apply only to your church.",
    "Global codes are shared across all churches.",
    "Windows seen by local admin now allow viewing all collection codes available to that church window:",
    "- church-specific codes for that local church", 
    "- global shared codes", 
    "By default, collection codes are created for every church from Collection_Codes.xlsx.",
    "The importer reads Collection_Codes.xlsx and materializes defaults per church during startup and when a new church is created.",
]

for line in lines:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.size = Pt(11)

doc.save(DOC_PATH)
print("Documentation.docx updated")
