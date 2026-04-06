#!/usr/bin/env python3
"""Update Documentation.docx with can_view_collection_codes section."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = 'Documentation.docx'

# Load document
doc = Document(doc_path)

# Add new paragraphs
doc.add_paragraph()  # Blank line

# Add heading
heading = doc.add_paragraph('LOCAL ADMIN: VIEWING COLLECTION CODES')
heading_format = heading.paragraph_format
heading_format.space_before = Pt(12)
heading_format.space_after = Pt(6)
heading.runs[0].bold = True
heading.runs[0].font.size = Pt(12)

# Add body content
body_text = """The new `can_view_collection_codes` permission allows local admins and other roles to view existing collection codes without necessarily having full management permissions.

Permission: can_view_collection_codes
• Applies to: system_admin, admin, treasurer, data_steward, uploader, viewer
• Effect: Users with this right can view all existing collection codes in their church on the Collections page
• Behavior: 
  - Collection codes will be visible in a read-only table
  - Only users with `can_manage_collections` can see the "Add New Code" section
  - Only users with `can_manage_collections` can edit or delete codes

Frontend Behavior:
• Collections menu button shows for users with either `can_manage_collections` OR `can_view_collection_codes`
• Collections page displays existing codes to viewers
• Management controls (Add, Edit, Delete) only appear for admins with `can_manage_collections` right

Database Schema:
• Table: role_policies
• New column: can_view_collection_codes (INTEGER, DEFAULT 0)
• Migration: ensure_role_policies_schema() adds column if missing and backfills for built-in roles

Operational Checklist:
✓ Backend restart applies migration to add can_view_collection_codes column
✓ Built-in roles (admin, treasurer, viewer, etc.) auto-populated with can_view_collection_codes=1
✓ Frontend renders Collection codes view for users with can_view_collection_codes right
✓ Edit/Delete buttons only appear for users with can_manage_collections right"""

for line in body_text.split('\n'):
    p = doc.add_paragraph(line)
    if line and not line.startswith('•') and not line.startswith('✓') and not line.startswith('-') and ':' not in line:
        p.paragraph_format.space_before = Pt(6)

# Save document
doc.save(doc_path)
print(f'Documentation updated successfully: {doc_path}')
